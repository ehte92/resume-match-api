"""
Resume parser service for extracting text and structure from PDF/DOCX files.
Extracts contact information, sections (experience, education, skills), and raw text.
"""

import re
from typing import Any, Dict

import docx
import pdfplumber
import PyPDF2


class ResumeParser:
    """Parse resume files (PDF/DOCX) and extract structured data"""

    def parse(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """
        Main entry point to parse resume

        Args:
            file_path: Path to resume file
            file_type: 'pdf' or 'docx'

        Returns:
            Dictionary with parsed resume data:
            {
                "raw_text": str,
                "email": str or None,
                "phone": str or None,
                "linkedin": str or None,
                "sections": {
                    "experience": [str],
                    "education": [str],
                    "skills": [str]
                }
            }

        Raises:
            ValueError: If file type is unsupported or parsing fails
        """
        # Extract text based on file type
        if file_type == "pdf":
            raw_text = self._extract_pdf(file_path)
        elif file_type == "docx":
            raw_text = self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        # Extract contact information
        contact_info = self._extract_contact_info(raw_text)

        # Identify sections
        sections = self._identify_sections(raw_text)

        return {
            "raw_text": raw_text,
            "email": contact_info.get("email"),
            "phone": contact_info.get("phone"),
            "linkedin": contact_info.get("linkedin"),
            "sections": sections,
        }

    def _extract_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file using PyPDF2 and pdfplumber

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text as string

        Raises:
            ValueError: If PDF extraction fails
        """
        text = ""

        try:
            # Try PyPDF2 first (faster)
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

            # If PyPDF2 didn't work well, try pdfplumber (better for complex layouts)
            if len(text.strip()) < 100:  # Very little text extracted
                text = ""
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"

        except Exception as e:
            raise ValueError(f"Error extracting PDF: {str(e)}")

        return text.strip()

    def _extract_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text as string

        Raises:
            ValueError: If DOCX extraction fails
        """
        try:
            doc = docx.Document(file_path)
            text = ""

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            return text.strip()

        except Exception as e:
            raise ValueError(f"Error extracting DOCX: {str(e)}")

    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        """
        Extract contact information using regex

        Args:
            text: Resume text to parse

        Returns:
            Dictionary with email, phone, linkedin (all optional)
        """
        contact = {}

        # Email regex
        email_pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
        email_match = re.search(email_pattern, text)
        if email_match:
            contact["email"] = email_match.group(0)

        # Phone regex (various formats)
        phone_pattern = r"(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}"
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            contact["phone"] = phone_match.group(0)

        # LinkedIn regex
        linkedin_pattern = r"linkedin\.com/in/[\w-]+"
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact["linkedin"] = linkedin_match.group(0)

        return contact

    def _identify_sections(self, text: str) -> Dict[str, list]:
        """
        Identify resume sections (Experience, Education, Skills)

        Args:
            text: Resume text to parse

        Returns:
            Dictionary with sections as keys and content as values
        """
        sections = {"experience": [], "education": [], "skills": []}

        # Split text into lines
        lines = text.split("\n")
        current_section = None
        section_content = []

        # Section heading patterns
        experience_patterns = [
            r"work\s+experience",
            r"professional\s+experience",
            r"experience",
            r"employment\s+history",
        ]
        education_patterns = [r"education", r"academic", r"qualifications"]
        skills_patterns = [r"skills", r"technical\s+skills", r"core\s+competencies"]

        for line in lines:
            line_lower = line.lower().strip()

            # Check if line is a section header
            if any(re.search(pattern, line_lower) for pattern in experience_patterns):
                # Save previous section content
                if current_section and section_content:
                    sections[current_section].append("\n".join(section_content))
                current_section = "experience"
                section_content = []

            elif any(re.search(pattern, line_lower) for pattern in education_patterns):
                # Save previous section content
                if current_section and section_content:
                    sections[current_section].append("\n".join(section_content))
                current_section = "education"
                section_content = []

            elif any(re.search(pattern, line_lower) for pattern in skills_patterns):
                # Save previous section content
                if current_section and section_content:
                    sections[current_section].append("\n".join(section_content))
                current_section = "skills"
                section_content = []

            # Add line to current section
            elif current_section and line.strip():
                section_content.append(line.strip())

        # Add last section
        if current_section and section_content:
            sections[current_section].append("\n".join(section_content))

        return sections
