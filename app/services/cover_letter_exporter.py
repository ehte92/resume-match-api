"""
Cover letter export service for generating PDF, DOCX, and TXT files.
Provides professional formatting for different export formats.
Handles both plain text and HTML-formatted cover letters.
"""

import io
import logging
import re
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.models.cover_letter import CoverLetter

logger = logging.getLogger(__name__)

# Font registration - register Noto Sans fonts for comprehensive Unicode support
FONTS_DIR = Path(__file__).parent.parent / "fonts"
FONTS_AVAILABLE = False

try:
    # Register regular Noto Sans fonts
    pdfmetrics.registerFont(TTFont("NotoSans", str(FONTS_DIR / "NotoSans-Regular.ttf")))
    pdfmetrics.registerFont(TTFont("NotoSans-Bold", str(FONTS_DIR / "NotoSans-Bold.ttf")))

    # Register Arabic fonts for RTL support
    pdfmetrics.registerFont(TTFont("NotoSansArabic", str(FONTS_DIR / "NotoSansArabic-Regular.ttf")))
    pdfmetrics.registerFont(
        TTFont("NotoSansArabic-Bold", str(FONTS_DIR / "NotoSansArabic-Bold.ttf"))
    )

    # Register font families for bold/italic support in HTML tags
    # This is CRITICAL for <b>, <i> tags to work in ReportLab Paragraphs
    pdfmetrics.registerFontFamily(
        'NotoSans',
        normal='NotoSans',
        bold='NotoSans-Bold',
        italic='NotoSans',  # No separate italic variant, use regular
        boldItalic='NotoSans-Bold'  # No separate bold-italic, use bold
    )

    pdfmetrics.registerFontFamily(
        'NotoSansArabic',
        normal='NotoSansArabic',
        bold='NotoSansArabic-Bold',
        italic='NotoSansArabic',
        boldItalic='NotoSansArabic-Bold'
    )

    FONTS_AVAILABLE = True
    logger.info("Successfully registered Noto Sans fonts and font families for PDF generation")
except Exception as e:
    logger.warning(f"Could not register Noto Sans fonts, falling back to Helvetica: {e}")
    FONTS_AVAILABLE = False


class CoverLetterExporter:
    """Export cover letters to various formats with professional styling."""

    @staticmethod
    def _contains_arabic(text: str) -> bool:
        """
        Check if text contains Arabic characters.

        Args:
            text: Text to check

        Returns:
            True if Arabic characters detected, False otherwise
        """
        arabic_range = range(0x0600, 0x06FF + 1)  # Arabic Unicode block
        return any(ord(char) in arabic_range for char in text)

    @staticmethod
    def _process_rtl_text(text: str) -> str:
        """
        Process Right-to-Left text (Arabic, Hebrew) for proper PDF rendering.

        Args:
            text: Text that may contain RTL content

        Returns:
            Processed text with proper RTL shaping
        """
        if not CoverLetterExporter._contains_arabic(text):
            return text

        try:
            from arabic_reshaper import reshape
            from bidi.algorithm import get_display

            # Reshape Arabic text (connect letters properly)
            reshaped_text = reshape(text)
            # Apply bidirectional algorithm for proper RTL display
            bidi_text = get_display(reshaped_text)
            return bidi_text
        except ImportError:
            logger.warning(
                "arabic-reshaper or python-bidi not available, Arabic text may not render correctly"
            )
            return text
        except Exception as e:
            logger.error(f"Error processing RTL text: {e}")
            return text

    @staticmethod
    def _is_html_content(text: str) -> bool:
        """
        Check if content contains HTML tags.

        Args:
            text: Text to check

        Returns:
            True if HTML tags detected, False otherwise
        """
        return bool(re.search(r"<[^>]+>", text))

    @staticmethod
    def _html_to_plain_text(html: str) -> str:
        """
        Convert HTML to plain text, preserving formatting.

        Args:
            html: HTML content

        Returns:
            Plain text with preserved structure
        """
        soup = BeautifulSoup(html, "html.parser")

        # Convert lists to bullet points
        for ul in soup.find_all("ul"):
            for li in ul.find_all("li"):
                li.string = f"• {li.get_text()}\n"

        for ol in soup.find_all("ol"):
            for idx, li in enumerate(ol.find_all("li"), 1):
                li.string = f"{idx}. {li.get_text()}\n"

        # Get text and clean up
        text = soup.get_text()
        # Remove excessive newlines
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @staticmethod
    def _normalize_text_for_pdf(text: str) -> str:
        """
        Normalize Unicode characters for PDF generation.
        Replaces special Unicode characters that Helvetica font doesn't support.

        Args:
            text: Text to normalize

        Returns:
            Normalized text with ASCII-safe characters
        """
        # Replace various dash/hyphen types with regular hyphen
        text = text.replace("\u2011", "-")  # Non-breaking hyphen
        text = text.replace("\u2013", "-")  # En dash
        text = text.replace("\u2014", "-")  # Em dash
        text = text.replace("\u2015", "-")  # Horizontal bar

        # Replace various quote types with straight quotes
        text = text.replace("\u2018", "'")  # Left single quote
        text = text.replace("\u2019", "'")  # Right single quote
        text = text.replace("\u201C", '"')  # Left double quote
        text = text.replace("\u201D", '"')  # Right double quote

        # Replace ellipsis
        text = text.replace("\u2026", "...")  # Horizontal ellipsis

        # Replace bullet variants
        text = text.replace("\u2022", "•")  # Bullet (keep this one as it works)
        text = text.replace("\u2023", "•")  # Triangular bullet

        return text

    @staticmethod
    def _html_to_reportlab_html(element) -> str:
        """
        Convert TipTap HTML element to ReportLab-compatible HTML.

        ReportLab's Paragraph supports: <b>, <i>, <u>, <br/>, <font>, <para>
        TipTap uses: <strong>, <em>, <u>, <p>

        This method:
        1. Converts <strong> → <b>, <em> → <i>
        2. Applies text normalization to text nodes only
        3. Preserves HTML structure for ReportLab rendering

        Args:
            element: BeautifulSoup element (paragraph or list item)

        Returns:
            HTML string compatible with ReportLab Paragraph
        """
        result = []

        # Process all children of the element
        for child in element.children:
            if child.name == "strong" or child.name == "b":
                # Bold text
                text_content = CoverLetterExporter._normalize_text_for_pdf(child.get_text())
                result.append(f"<b>{text_content}</b>")
            elif child.name == "em" or child.name == "i":
                # Italic text
                text_content = CoverLetterExporter._normalize_text_for_pdf(child.get_text())
                result.append(f"<i>{text_content}</i>")
            elif child.name == "u":
                # Underline text
                text_content = CoverLetterExporter._normalize_text_for_pdf(child.get_text())
                result.append(f"<u>{text_content}</u>")
            elif child.name == "br":
                # Line break
                result.append("<br/>")
            elif child.name is None:
                # Text node - normalize and escape HTML special chars
                text = str(child)
                text = CoverLetterExporter._normalize_text_for_pdf(text)
                # Escape HTML special characters
                text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                result.append(text)
            else:
                # Other tags - extract text and normalize
                text_content = CoverLetterExporter._normalize_text_for_pdf(child.get_text())
                # Escape HTML special characters
                text_content = text_content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                result.append(text_content)

        return "".join(result)

    @staticmethod
    def export_to_txt(cover_letter: CoverLetter) -> bytes:
        """
        Export cover letter as plain text file.
        Handles both HTML and plain text content.

        Args:
            cover_letter: CoverLetter model instance

        Returns:
            UTF-8 encoded text bytes
        """
        # Build header with job details
        header_lines = []
        if cover_letter.job_title:
            header_lines.append(f"Position: {cover_letter.job_title}")
        if cover_letter.company_name:
            header_lines.append(f"Company: {cover_letter.company_name}")
        header_lines.append(f"Generated: {cover_letter.created_at.strftime('%B %d, %Y')}")
        header_lines.append("")  # Empty line

        # Convert HTML to plain text if needed
        letter_text = cover_letter.cover_letter_text
        if CoverLetterExporter._is_html_content(letter_text):
            letter_text = CoverLetterExporter._html_to_plain_text(letter_text)

        # Assemble full content
        content = "\n".join(header_lines) + "\n" + letter_text

        return content.encode("utf-8")

    @staticmethod
    def export_to_pdf(cover_letter: CoverLetter, include_metadata: bool = False) -> bytes:
        """
        Export cover letter as professionally formatted PDF.

        Args:
            cover_letter: CoverLetter model instance
            include_metadata: Whether to include generation metadata in footer

        Returns:
            PDF file as bytes
        """
        buffer = io.BytesIO()

        # Create PDF with standard letter size and margins
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=inch,
            leftMargin=inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )

        # Build story (content elements)
        story = []
        styles = getSampleStyleSheet()

        # Determine font names based on availability
        base_font = "NotoSans" if FONTS_AVAILABLE else "Helvetica"
        bold_font = "NotoSans-Bold" if FONTS_AVAILABLE else "Helvetica-Bold"

        # Create custom styles
        heading_style = ParagraphStyle(
            "CustomHeading",
            parent=styles["Heading2"],
            fontSize=14,
            textColor=colors.black,
            spaceAfter=6,
            fontName=bold_font,
        )

        subheading_style = ParagraphStyle(
            "CustomSubheading",
            parent=styles["Normal"],
            fontSize=10,
            textColor=colors.grey,
            spaceAfter=12,
            fontName=base_font,
        )

        body_style = ParagraphStyle(
            "CustomBody",
            parent=styles["Normal"],
            fontSize=11,
            leading=16,
            textColor=colors.black,
            fontName=base_font,
            alignment=TA_LEFT,  # Use ReportLab alignment constant
        )

        # Add header with job details
        if cover_letter.job_title:
            normalized_title = CoverLetterExporter._normalize_text_for_pdf(cover_letter.job_title)
            # Process RTL text if needed
            processed_title = CoverLetterExporter._process_rtl_text(normalized_title)
            story.append(Paragraph(processed_title, heading_style))
        if cover_letter.company_name:
            normalized_company = CoverLetterExporter._normalize_text_for_pdf(
                cover_letter.company_name
            )
            # Process RTL text if needed
            processed_company = CoverLetterExporter._process_rtl_text(normalized_company)
            story.append(Paragraph(processed_company, subheading_style))
        else:
            story.append(Spacer(1, 12))

        story.append(Spacer(1, 0.2 * inch))

        # Add cover letter text with formatting
        letter_text = cover_letter.cover_letter_text

        if CoverLetterExporter._is_html_content(letter_text):
            # Parse HTML and convert to ReportLab format
            soup = BeautifulSoup(letter_text, "html.parser")

            # Process each paragraph element
            for element in soup.find_all(["p", "ul", "ol"]):
                if element.name == "p":
                    # Handle paragraph with inline formatting
                    # Convert TipTap HTML to ReportLab-compatible HTML
                    text = CoverLetterExporter._html_to_reportlab_html(element)
                    # Process RTL text if needed (after HTML conversion)
                    text = CoverLetterExporter._process_rtl_text(text)
                    # ReportLab will render <b>, <i>, <u>, <br/> tags
                    para = Paragraph(text, body_style)
                    story.append(para)
                    story.append(Spacer(1, 12))
                elif element.name in ["ul", "ol"]:
                    # Handle lists with inline formatting support
                    list_items = element.find_all("li", recursive=False)
                    for idx, li in enumerate(list_items):
                        bullet = "•" if element.name == "ul" else f"{idx + 1}."
                        # Convert list item HTML to ReportLab format
                        li_text = CoverLetterExporter._html_to_reportlab_html(li)
                        # Add bullet/number and process RTL text
                        text = f"{bullet} {li_text}"
                        text = CoverLetterExporter._process_rtl_text(text)
                        para = Paragraph(text, body_style)
                        story.append(para)
                        story.append(Spacer(1, 6))
                    story.append(Spacer(1, 6))
        else:
            # Plain text - split on double newlines for paragraphs
            paragraphs = letter_text.split("\n\n")
            for para_text in paragraphs:
                if para_text.strip():
                    # Normalize Unicode characters first
                    normalized_text = CoverLetterExporter._normalize_text_for_pdf(para_text.strip())
                    # Process RTL text if needed
                    processed_text = CoverLetterExporter._process_rtl_text(normalized_text)
                    # Escape HTML special characters and replace line breaks
                    safe_text = (
                        processed_text.replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;")
                    )
                    safe_text = safe_text.replace("\n", "<br/>")
                    para = Paragraph(safe_text, body_style)
                    story.append(para)
                    story.append(Spacer(1, 12))

        # Add metadata footer if requested
        if include_metadata:
            story.append(Spacer(1, 0.3 * inch))
            # Use base font for metadata (Noto Sans doesn't have oblique variant)
            metadata_font = base_font if FONTS_AVAILABLE else "Helvetica-Oblique"
            metadata_style = ParagraphStyle(
                "Metadata",
                parent=styles["Normal"],
                fontSize=8,
                textColor=colors.grey,
                fontName=metadata_font,
            )
            metadata_text = CoverLetterExporter._normalize_text_for_pdf(
                f"Generated on {cover_letter.created_at.strftime('%B %d, %Y')} | "
                f"Tone: {cover_letter.tone.capitalize()} | "
                f"Length: {cover_letter.length.capitalize()} | "
                f"Words: {cover_letter.word_count}"
            )
            # Process RTL text if needed
            metadata_text = CoverLetterExporter._process_rtl_text(metadata_text)
            story.append(Paragraph(metadata_text, metadata_style))

        # Build PDF
        doc.build(story)
        pdf_bytes = buffer.getvalue()
        buffer.close()

        logger.info(f"Generated PDF export for cover letter {cover_letter.id}")
        return pdf_bytes

    @staticmethod
    def export_to_docx(cover_letter: CoverLetter) -> bytes:
        """
        Export cover letter as Microsoft Word document (.docx).

        Args:
            cover_letter: CoverLetter model instance

        Returns:
            DOCX file as bytes
        """
        doc = Document()

        # Set document margins (1 inch all sides)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Add header with job details
        if cover_letter.job_title:
            heading = doc.add_heading(cover_letter.job_title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if cover_letter.company_name:
            company_para = doc.add_paragraph(cover_letter.company_name)
            company_para.runs[0].font.size = Pt(11)
            company_para.runs[0].font.color.rgb = None  # Default color

        # Add spacing
        doc.add_paragraph()

        # Add cover letter text with formatting
        letter_text = cover_letter.cover_letter_text

        if CoverLetterExporter._is_html_content(letter_text):
            # Parse HTML and apply formatting to DOCX
            soup = BeautifulSoup(letter_text, "html.parser")

            for element in soup.find_all(["p", "ul", "ol"]):
                if element.name == "p":
                    # Create paragraph
                    para = doc.add_paragraph()
                    # Parse inline formatting (bold, italic, underline)
                    for child in element.children:
                        if child.name == "strong" or child.name == "b":
                            run = para.add_run(child.get_text())
                            run.bold = True
                        elif child.name == "em" or child.name == "i":
                            run = para.add_run(child.get_text())
                            run.italic = True
                        elif child.name == "u":
                            run = para.add_run(child.get_text())
                            run.underline = True
                        elif child.name == "br":
                            para.add_run("\n")
                        elif child.name is None:  # Text node
                            para.add_run(str(child))
                        else:
                            para.add_run(child.get_text())

                    # Apply formatting to all runs
                    for run in para.runs:
                        if not run.font.name:
                            run.font.name = "Calibri"
                        if not run.font.size:
                            run.font.size = Pt(11)
                    para.paragraph_format.line_spacing = 1.15
                    para.paragraph_format.space_after = Pt(10)

                elif element.name == "ul":
                    # Bullet list
                    for li in element.find_all("li", recursive=False):
                        para = doc.add_paragraph(li.get_text(), style="List Bullet")
                        for run in para.runs:
                            run.font.name = "Calibri"
                            run.font.size = Pt(11)

                elif element.name == "ol":
                    # Numbered list
                    for li in element.find_all("li", recursive=False):
                        para = doc.add_paragraph(li.get_text(), style="List Number")
                        for run in para.runs:
                            run.font.name = "Calibri"
                            run.font.size = Pt(11)
        else:
            # Plain text - split on double newlines for paragraphs
            paragraphs = letter_text.split("\n\n")
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text.strip())
                    # Set font and spacing
                    for run in para.runs:
                        run.font.name = "Calibri"
                        run.font.size = Pt(11)
                    para.paragraph_format.line_spacing = 1.15
                    para.paragraph_format.space_after = Pt(10)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()

        logger.info(f"Generated DOCX export for cover letter {cover_letter.id}")
        return docx_bytes

    @staticmethod
    def get_filename(cover_letter: CoverLetter, export_format: str) -> str:
        """
        Generate appropriate filename for exported cover letter.

        Args:
            cover_letter: CoverLetter model instance
            export_format: File format (pdf, docx, txt)

        Returns:
            Filename string (e.g., "Google-Software-Engineer-cover-letter.pdf")
        """
        parts = []

        # Add company name
        if cover_letter.company_name:
            # Remove non-ASCII characters for HTTP header compatibility
            safe_company = "".join(c if ord(c) < 128 else "" for c in cover_letter.company_name)
            safe_company = safe_company.strip().replace(" ", "-")
            if safe_company:  # Only add if there's something left after sanitization
                parts.append(safe_company)

        # Add job title
        if cover_letter.job_title:
            # Remove non-ASCII characters for HTTP header compatibility
            safe_title = "".join(c if ord(c) < 128 else "" for c in cover_letter.job_title)
            safe_title = safe_title.strip().replace(" ", "-")
            if safe_title:  # Only add if there's something left after sanitization
                parts.append(safe_title)

        # Default if no details provided
        if not parts:
            parts.append("cover-letter")

        # Add format extension
        filename = "-".join(parts) + f".{export_format}"

        # Clean filename (remove invalid characters)
        invalid_chars = ["/", "\\", ":", "*", "?", '"', "<", ">", "|"]
        for char in invalid_chars:
            filename = filename.replace(char, "-")

        logger.info(f"Generated filename: {filename}")
        return filename
